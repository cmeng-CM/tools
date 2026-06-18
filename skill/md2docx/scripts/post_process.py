#!/usr/bin/env python3
"""
Post-process a pandoc-generated .docx for professional formatting.

Fixes that pandoc doesn't handle (or handles poorly):
  1. Table borders + header row formatting
  2. Image sizing (constrain to page width)
  3. Hyperlink style verification
  4. CJK font consistency on special styles

Usage:
    python3 scripts/post_process.py <output.docx> [--backup]

All operations are idempotent — safe to re-run on the same file.
"""

import os
import sys
import shutil

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Emu
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.style import WD_STYLE_TYPE


# ── Helpers ─────────────────────────────────────────────────────────

def _set_cell_border(cell, sides, color='808080', sz=4, style='single'):
    """Add borders to a table cell. sides: list of 'top'/'bottom'/'left'/'right'."""
    tc = cell._tc
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = OxmlElement('w:tcPr')
        tc.insert(0, tcPr)
    borders = tcPr.find(qn('w:tcBorders'))
    if borders is None:
        borders = OxmlElement('w:tcBorders')
        tcPr.append(borders)
    for side in sides:
        existing = borders.find(qn(f'w:{side}'))
        if existing is not None:
            borders.remove(existing)
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), style)
        b.set(qn('w:sz'), str(sz))
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), color)
        borders.append(b)


def _set_cell_shading(cell, fill):
    """Set background shading on a cell."""
    tc = cell._tc
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = OxmlElement('w:tcPr')
        tc.insert(0, tcPr)
    existing = tcPr.find(qn('w:shd'))
    if existing is not None:
        tcPr.remove(existing)
    s = OxmlElement('w:shd')
    s.set(qn('w:val'), 'clear')
    s.set(qn('w:color'), 'auto')
    s.set(qn('w:fill'), fill)
    tcPr.append(s)


def _get_elem(parent, tag):
    """Get or create an XML child element."""
    el = parent.find(qn(tag))
    if el is None:
        el = OxmlElement(tag)
        parent.insert(0, el)
    return el


# ── Operation 1: Table Formatting ───────────────────────────────────

def format_tables(doc):
    """Add borders to all tables, bold + shade header rows."""
    tables = doc.tables
    if not tables:
        return 0

    header_bg = 'E7E6E6'       # lt2 from theme
    border_color = '808080'

    for table in tables:
        rows = table.rows
        if not rows:
            continue

        # Bold header row and shade
        for cell in rows[0].cells:
            _set_cell_shading(cell, header_bg)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.bold = True

        # Add borders to all cells
        all_sides = ['top', 'bottom', 'left', 'right']
        for row in rows:
            for cell in row.cells:
                _set_cell_border(cell, all_sides, border_color, sz=4)

    print(f'  Tables formatted: {len(tables)}')
    return len(tables)


# ── Operation 2: Image Sizing ───────────────────────────────────────

def resize_images(doc):
    """Constrain inline images to page content width."""
    # Calculate page content width
    section = doc.sections[0]
    page_w = section.page_width
    left_m = section.left_margin
    right_m = section.right_margin

    # If margins not set, use defaults
    if left_m is None:
        left_m = Cm(2.54)
    if right_m is None:
        right_m = Cm(2.54)

    content_w = page_w - left_m - right_m
    max_w = int(content_w * 0.9)  # 90% of content width

    nsmap = {
        'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
        'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
    }

    count = 0
    for p in doc.paragraphs:
        drawings = p._p.findall(qn('w:r') + '/' + qn('w:drawing'))
        if not drawings:
            # Check inside runs
            drawings = []
            for r in p._p.findall(qn('w:r')):
                drawings.extend(r.findall(qn('w:drawing')))

        for drawing in drawings:
            # Find extent (current size in EMU)
            extents = drawing.iter(qn('wp:extent'))
            extent_list = list(extents)
            if not extents:
                # Try a:ext
                extent_list = list(drawing.iter(qn('a:ext')))

            if not extent_list:
                continue

            ext = extent_list[0]
            cx = int(ext.get('cx', '0'))
            cy = int(ext.get('cy', '0'))

            if cx > max_w and cx > 0:
                ratio = max_w / cx
                new_cx = max_w
                new_cy = int(cy * ratio)
                ext.set('cx', str(new_cx))
                ext.set('cy', str(new_cy))
                count += 1

    if count:
        print(f'  Images resized: {count}')
    return count


# ── Operation 3: Hyperlink Style ────────────────────────────────────

def ensure_hyperlink_style(doc):
    """Make sure Hyperlink character style exists with proper formatting."""
    try:
        doc.styles['Hyperlink']
        return False  # already exists
    except KeyError:
        s = doc.styles.add_style('Hyperlink', WD_STYLE_TYPE.CHARACTER)
        s.hidden = False
        s.font.color.rgb = RGBColor(5, 99, 193)  # #0563C1
        s.font.underline = True
        print('  Hyperlink style added')
        return True


# ── Operation 4: List Spacing ───────────────────────────────────────

def fix_list_spacing(doc):
    """Tighten spacing on list paragraphs using Compact style."""
    count = 0
    for p in doc.paragraphs:
        style_name = p.style.name if p.style else ''
        if style_name == 'Compact':
            pf = p.paragraph_format
            if pf.space_after is None or pf.space_after > Pt(4):
                pf.space_after = Pt(2)
                count += 1
    if count:
        print(f'  List paragraphs tightened: {count}')
    return count


# ── Operation 5: Reposition TOC after document title ─────────────────

def reposition_toc_after_title(doc):
    """Move the TOC (sdt) from the very beginning to after the first heading
    (any level — the document title), and enable auto-update on document open
    so the TOC populates without requiring the user to manually right-click →
    Update Field.
    """
    body = doc.element.body

    # Find the TOC sdt element
    toc_sdt = None
    for sdt in body.findall(qn('w:sdt')):
        doc_part = sdt.find('.//' + qn('w:docPartGallery'))
        if doc_part is not None and doc_part.get(qn('w:val')) == 'Table of Contents':
            toc_sdt = sdt
            break

    moved = False

    if toc_sdt is not None:
        # --- Reposition TOC after the first heading (any level) ---
        children = list(body)

        # Find the first heading — pandoc maps # → '1'/Heading1, ## → '2'/Heading2, etc.
        heading_idx = None
        for i, child in enumerate(children):
            if child.tag == qn('w:p'):
                pPr = child.find(qn('w:pPr'))
                if pPr is not None:
                    pStyle = pPr.find(qn('w:pStyle'))
                    if pStyle is not None:
                        style_id = (pStyle.get(qn('w:val')) or '').lower()
                        if style_id.isdigit() or style_id.startswith('heading'):
                            heading_idx = i
                            break

        if heading_idx is not None:
            toc_idx = children.index(toc_sdt)
            if toc_idx < heading_idx:
                body.remove(toc_sdt)
                body.insert(heading_idx + 1, toc_sdt)
                moved = True

    # --- Enable auto-update of fields on document open ---
    # This tells Word to update the TOC automatically when the document
    # is first opened, so users don't need to right-click → Update Field.
    settings = doc.settings.element
    update_fields = settings.find(qn('w:updateFields'))
    if update_fields is None:
        update_fields = OxmlElement('w:updateFields')
        settings.append(update_fields)
    update_fields.set(qn('w:val'), 'true')

    return moved


# ── Main ────────────────────────────────────────────────────────────

def post_process(docx_path):
    """Run all post-processing operations on a docx file."""
    if not os.path.exists(docx_path):
        print(f'ERROR: file not found: {docx_path}')
        sys.exit(1)

    doc = Document(docx_path)

    print(f'Processing: {docx_path}')
    toc_moved = reposition_toc_after_title(doc)
    n_tables = format_tables(doc)
    n_images = resize_images(doc)
    hyper_added = ensure_hyperlink_style(doc)
    n_lists = fix_list_spacing(doc)

    doc.save(docx_path)

    toc_status = 'moved' if toc_moved else 'none'
    print(f'Done — toc:{toc_status} tables:{n_tables} images:{n_images} '
          f'hyperlink:{"added" if hyper_added else "ok"} lists:{n_lists}')


if __name__ == '__main__':
    import argparse
    p = argparse.ArgumentParser(
        description='Post-process pandoc-generated .docx')
    p.add_argument('docx', help='Path to the .docx file')
    args = p.parse_args()
    post_process(args.docx)
